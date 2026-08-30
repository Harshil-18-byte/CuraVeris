import io
from datetime import datetime, timezone
from decimal import Decimal
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Response, status
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.orm import selectinload
from app.db.database import get_db
from app.db.models import Bill, DisputeLetter, User
from app.models.schemas import DisputeLetterRequest, DisputeLetterResponse
from app.services.dispute_service import dispute_service
from app.api.auth import get_optional_user
from app.core.security_hardening import SecurityHardeningEngine

router = APIRouter(prefix="/reports", tags=["Reports & Disputes"])



@router.post("/dispute-letter", response_model=DisputeLetterResponse)
async def generate_dispute_letter(
    req: DisputeLetterRequest,
    current_user: Optional[User] = Depends(get_optional_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Generate formal legal petition (Hospital Grievance, NPPA, IRDAI, Consumer Court).
    """
    result = await db.execute(
        select(Bill).where(Bill.id == req.bill_id).options(selectinload(Bill.items))
    )
    bill = result.scalars().first()
    if not bill:
        raise HTTPException(status_code=404, detail="Bill not found")

    # IDOR check: if bill has an owner and caller is a PATIENT, verify ownership
    if bill.user_id and current_user:
        is_staff = current_user.role in ("HOSPITAL_ADMIN", "HOSPITAL_FINANCE", "HOSPITAL_BILLING", "HOSPITAL_AUDITOR", "TPA_REVIEWER", "TPA_ADMIN", "INSURER_REVIEWER", "INSURER_ADMIN", "PLATFORM_ADMIN")
        if not is_staff and str(bill.user_id) != str(current_user.id):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access forbidden: you do not have permission to generate dispute letters for this bill."
            )


    bill_data = {
        "bill_id": bill.id,
        "hospital_name": bill.hospital_name,
        "city": bill.city,
        "total_billed": bill.total_billed,
        "total_overcharge": bill.total_overcharge,
        "line_items": [
            {
                "id": i.id,
                "raw_text": i.raw_text,
                "charged_rate": i.charged_rate,
                "quantity": i.quantity,
                "overcharge_amount": i.overcharge_amount,
                "is_flagged": i.is_flagged,
                "legal_citation": i.legal_citation,
                "mrp": i.mrp,
                "nppa_ceiling": i.nppa_ceiling,
                "cghs_rate": i.cghs_rate
            }
            for i in bill.items
        ]
    }

    forum_type_str = str(req.forum_type or "HOSPITAL_GRIEVANCE")
    letter_data = dispute_service.generate_letter(
        bill_data=bill_data,
        forum_type=forum_type_str,
        patient_name=req.patient_name or "Aggrieved Patient",
        patient_address=req.patient_address or "India",
        patient_phone=req.patient_phone or "+91-9876543210"
    )

    # Save into DB
    letter_obj = DisputeLetter(
        bill_id=str(bill.id),
        forum_type=str(letter_data["forum_type"]),
        target_authority=str(letter_data["target_authority"]),
        letter_title=str(letter_data["letter_title"]),
        letter_body=str(letter_data["letter_body"]),
        statutory_citations=list(letter_data["statutory_citations"]),
        total_disputed_amount=float(letter_data["total_disputed_amount"])
    )
    db.add(letter_obj)
    await db.commit()
    await db.refresh(letter_obj)

    from decimal import Decimal
    return DisputeLetterResponse(
        letter_id=str(letter_obj.id),
        bill_id=str(letter_obj.bill_id),
        target_authority=str(letter_obj.target_authority),
        letter_title=str(letter_obj.letter_title) if letter_obj.letter_title else None,
        letter_body=str(letter_obj.letter_body) if letter_obj.letter_body else None,
        statutory_citations=[str(c) for c in (letter_data.get("statutory_citations") or [])],
        total_disputed_amount=Decimal(str(letter_data.get("total_disputed_amount", 0.0))),
        created_at=datetime.now(timezone.utc)
    )




@router.get("/docx/{letter_id}")
async def download_dispute_docx(
    letter_id: str,
    current_user: Optional[User] = Depends(get_optional_user),
    db: AsyncSession = Depends(get_db)
):
    """Generate and download petition as a formatted Word Document (.docx)."""
    result = await db.execute(select(DisputeLetter).where(DisputeLetter.id == letter_id))
    letter = result.scalars().first()
    if not letter:
        raise HTTPException(status_code=404, detail="Dispute letter not found")

    # IDOR check: verify bill ownership if bound
    if letter.bill_id and current_user:
        bill_res = await db.execute(select(Bill).where(Bill.id == letter.bill_id))
        bill = bill_res.scalars().first()
        if bill and bill.user_id:
            is_staff = current_user.role in ("HOSPITAL_ADMIN", "HOSPITAL_FINANCE", "HOSPITAL_BILLING", "HOSPITAL_AUDITOR", "TPA_REVIEWER", "TPA_ADMIN", "INSURER_REVIEWER", "INSURER_ADMIN", "PLATFORM_ADMIN")
            if not is_staff and str(bill.user_id) != str(current_user.id):
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Access forbidden: you do not have permission to download this dispute letter."
                )

    try:

        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        
        doc = Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(letter.letter_title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.color.rgb = RGBColor(30, 58, 95)
        
        doc.add_paragraph()
        
        # Body lines
        for line in letter.letter_body.splitlines():
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"dispute_{letter.forum_type}_{letter.id[:6]}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to render docx: {e}")


@router.post("/emergency-detention-notice")
async def generate_emergency_detention_notice(
    hospital_name: str,
    patient_name: str,
    attendant_name: str,
    attendant_phone: str,
    disputed_amount: float,
    city: str = "India"
):
    """
    Emergency High Court Cease-and-Desist Notice against Illegal Patient Detention.
    Generates an immediate formal requisition citing:
    - Bombay High Court in 'Association of Medical Consultants vs Union of India'
    - Bharatiya Nyaya Sanhita 2023 Sec 127 (IPC 340/342 Wrongful Confinement)
    - Article 21 of the Constitution of India (Personal Liberty)
    """
    notice = dispute_service.generate_emergency_detention_notice(
        hospital_name=hospital_name,
        patient_name=patient_name,
        attendant_name=attendant_name,
        attendant_phone=attendant_phone,
        disputed_amount=disputed_amount,
        city=city
    )
    return notice

